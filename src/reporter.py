"""
Daily reporter — Hitman v2 methodology.
Runs at 4:30 PM ET weekdays.
Compiles daily P&L, trade summary, and budget status.
Saves HTML report to reports/ and emails it.
Also updates earnings_calls_log.json.
"""
from datetime import date, timedelta
from pathlib import Path
from typing import Dict, Any, List, Optional

import yfinance as yf

from utils import (
    get_logger, load_settings, load_json, save_json,
    get_trading_client, send_email, now_et, today_et, is_weekday,
    REPORTS_DIR,
)

DATA_DIR = Path(__file__).resolve().parent.parent / "data"
EARNINGS_DOCX = DATA_DIR / "earnings_calls_log.docx"

log = get_logger("reporter")


def get_today_trades(trades_history: List[Dict]) -> List[Dict]:
    today = today_et().isoformat()
    return [t for t in trades_history if t.get("entry_date") == today]


def get_today_closed(trades_history: List[Dict]) -> List[Dict]:
    today = today_et().isoformat()
    return [
        t for t in trades_history
        if t.get("exit_date") == today and t.get("status") == "closed"
    ]


def get_all_closed(trades_history: List[Dict]) -> List[Dict]:
    return [t for t in trades_history if t.get("status") == "closed"]


def calc_win_rate(closed_trades: List[Dict]) -> float:
    if not closed_trades:
        return 0.0
    wins = sum(1 for t in closed_trades if (t.get("pnl") or 0) > 0)
    return wins / len(closed_trades) * 100




def build_text_report(
    today_trades: List[Dict],
    today_closed: List[Dict],
    all_closed: List[Dict],
    positions: Dict,
    budget: Dict,
) -> str:
    today          = today_et().isoformat()
    daily_pnl      = positions.get("daily_pnl", 0)
    daily_pnl_pct  = positions.get("daily_pnl_pct", 0)
    open_positions = positions.get("positions", [])
    win_rate       = calc_win_rate(all_closed)
    total_pnl      = sum(t.get("pnl", 0) or 0 for t in all_closed)

    lines = [
        f"HITMAN v2 DAILY REPORT — {today}",
        "=" * 60,
        f"Daily P&L:       ${daily_pnl:+,.2f} ({daily_pnl_pct:+.1f}%)",
        f"Total P&L:       ${total_pnl:+,.2f}",
        f"Win Rate:        {win_rate:.0f}% ({len(all_closed)} closed trades)",
        f"Open Positions:  {len(open_positions)}",
        "",
        "BUDGET STATUS",
        "-" * 40,
        f"  Deployed (open):  ${budget.get('total_deployed', 0):,.0f}",
        f"  Total cap:        ${budget.get('total_cap', 1_000_000):,.0f}",
        f"  Available:        ${budget.get('available_total', 1_000_000):,.0f}",
        "",
    ]

    if today_trades:
        lines.append("OPTIONS TRADES OPENED TODAY")
        lines.append("-" * 40)
        for t in today_trades:
            contract = t.get("long_contract", "N/A")
            lines.append(
                f"  {t['symbol']:6s} | {t.get('trade_type','?'):6s} | "
                f"{t['qty']} contracts | contract={contract} | "
                f"cost=${t['position_size']:,.0f} | earnings={t['earnings_date']}"
            )
        lines.append("")

    if today_closed:
        lines.append("OPTIONS TRADES CLOSED TODAY")
        lines.append("-" * 40)
        for t in today_closed:
            outcome = "WIN" if (t.get("pnl") or 0) > 0 else "LOSS"
            lines.append(
                f"  {t['symbol']:6s} | {outcome} | {t.get('trade_type','?'):6s} | "
                f"cost=${t.get('position_size',0):,.0f} → "
                f"exit=${t.get('exit_value', t.get('exit_price', 0)):,.2f} | "
                f"P&L: ${t.get('pnl', 0):+,.2f} ({t.get('pnl_pct', 0):+.1f}%)"
            )
        lines.append("")

    if open_positions:
        lines.append("OPEN POSITIONS")
        lines.append("-" * 40)
        for p in open_positions:
            pnl_pct = p.get("unrealized_plpc", 0) * 100
            lines.append(
                f"  {p['symbol']:30s} | qty={p['qty']:.0f} | "
                f"entry=${p['avg_entry_price']:.4f} | now=${p['current_price']:.4f} | "
                f"P&L: ${p['unrealized_pl']:+,.2f} ({pnl_pct:+.1f}%)"
            )
        lines.append("")

    cb   = positions.get("circuit_breaker_triggered")
    halt = budget.get("halted")
    if cb or halt:
        lines.append("ALERTS")
        lines.append("-" * 40)
        if cb:
            lines.append(f"  CIRCUIT BREAKER: {positions.get('circuit_breaker_reason')}")
        if halt:
            lines.append(f"  TRADING HALTED: {budget.get('halt_reason')}")

    return "\n".join(lines)


def build_html_report(text_report: str, today_closed: List[Dict]) -> str:
    wins = sum(1 for t in today_closed if (t.get("pnl") or 0) > 0)
    losses = len(today_closed) - wins
    return f"""<!DOCTYPE html>
<html>
<head>
<style>
body {{ font-family: monospace; background: #0d1117; color: #c9d1d9; padding: 20px; }}
h1 {{ color: #58a6ff; }}
.win {{ color: #3fb950; }}
.loss {{ color: #f85149; }}
pre {{ background: #161b22; padding: 15px; border-radius: 6px; overflow-x: auto; }}
</style>
</head>
<body>
<h1>Daily Performance Report — {today_et().isoformat()}</h1>
<p>Today: <span class="win">{wins} wins</span> / <span class="loss">{losses} losses</span></p>
<pre>{text_report}</pre>
</body>
</html>"""


def fetch_earnings_result(symbol: str, expected_date: str) -> Dict[str, Any]:
    """Try to get actual EPS beat/miss from yfinance for the given earnings date."""
    result = {"beat": None, "eps_estimate": None, "eps_reported": None, "beat_pct": None}
    try:
        ticker = yf.Ticker(symbol)
        edf = ticker.earnings_dates
        if edf is None or edf.empty:
            return result
        # Look within 5 days of the expected earnings date
        expected = date.fromisoformat(expected_date)
        edf.index = edf.index.tz_localize(None) if edf.index.tzinfo else edf.index
        for ts, row in edf.iterrows():
            row_date = ts.date() if hasattr(ts, 'date') else ts
            if abs((row_date - expected).days) <= 5:
                est = row.get("EPS Estimate")
                rep = row.get("Reported EPS")
                if est is not None and rep is not None and str(est) != 'nan' and str(rep) != 'nan':
                    est, rep = float(est), float(rep)
                    beat = rep > est
                    beat_pct = ((rep - est) / abs(est) * 100) if est != 0 else 0
                    result.update({"beat": beat, "eps_estimate": round(est, 3),
                                   "eps_reported": round(rep, 3), "beat_pct": round(beat_pct, 1)})
                    return result
    except Exception:
        pass
    return result


def fetch_post_earnings_news(symbol: str, earnings_date: str) -> Optional[str]:
    """
    Pull news headlines published within 3 days after earnings_date via yfinance.
    Reads the top 5 articles, extracts the dominant theme, and returns a single
    plain-English sentence explaining why the stock moved against expectations.
    Returns None if no relevant articles are found.
    """
    try:
        from datetime import timedelta as _td
        earnings_d  = date.fromisoformat(earnings_date)
        window_end  = earnings_d + _td(days=3)

        ticker   = yf.Ticker(symbol)
        articles = ticker.news or []

        # Collect titles from articles published within the post-earnings window
        titles = []
        for a in articles:
            pub_ts = a.get("providerPublishTime") or 0
            if not pub_ts:
                continue
            try:
                pub_date = date.fromtimestamp(int(pub_ts))
            except Exception:
                continue
            if earnings_d <= pub_date <= window_end:
                title = (a.get("title") or "").strip()
                if title:
                    titles.append(title)

        if not titles:
            return None

        combined = " ".join(titles[:5]).lower()

        # Theme detection — most specific patterns first
        if any(w in combined for w in ["guidance cut", "lowered guidance", "cut guidance",
                                        "guidance disappoint", "weak guidance", "slashed guidance"]):
            return (f"Post-earnings coverage flagged a guidance cut as the primary driver "
                    f"of the selloff, which overshadowed the EPS beat.")

        if any(w in combined for w in ["revenue miss", "missed revenue", "sales miss",
                                        "top-line miss", "revenue fell short", "revenue disappoint"]):
            return (f"Coverage pointed to a revenue miss — investors focused on the weaker "
                    f"top-line result rather than the headline EPS beat.")

        if any(w in combined for w in ["guidance", "outlook", "forecast", "forward guidance",
                                        "next quarter", "full-year"]):
            return (f"Post-earnings articles highlighted concerns around forward guidance, "
                    f"which the market weighted more heavily than the earnings beat.")

        if any(w in combined for w in ["tariff", "trade war", "macro", "inflation",
                                        "interest rate", "recession", "economy", "fed "]):
            return (f"Coverage pointed to macro headwinds — tariff or economic uncertainty "
                    f"drove the selloff despite the earnings beat.")

        if any(w in combined for w in ["margin", "gross margin", "operating margin",
                                        "profit margin", "margin compress"]):
            return (f"Articles flagged margin compression as the key concern weighing on "
                    f"the stock, despite the headline EPS beat.")

        if any(w in combined for w in ["competition", "market share", "competitive pressure",
                                        "rival", "losing share"]):
            return (f"Post-earnings coverage cited competitive pressure as the reason "
                    f"investors sold into the earnings beat.")

        if any(w in combined for w in ["valuation", "expensive", "overvalued",
                                        "priced in", "priced for perfection", "sell the news"]):
            return (f"Coverage suggested the beat was already priced in, leading to a "
                    f"classic sell-the-news reaction.")

        if any(w in combined for w in ["disappoint", "fell short", "concern", "worry",
                                        "below expect", "not enough"]):
            return (f"Despite the EPS beat, post-earnings coverage reflected broader "
                    f"investor concerns that drove the negative price reaction.")

        # Last resort — surface the most earnings-relevant headline directly
        best = next(
            (t for t in titles if any(k in t.lower() for k in
             ["beat", "earnings", "results", "guidance", "revenue", "profit"])),
            titles[0],
        )
        short = best[:120] + ("..." if len(best) > 120 else "")
        return (f'Post-earnings research found: "{short}" — suggesting factors beyond '
                f"the EPS beat drove the adverse reaction.")

    except Exception:
        return None


def diagnose_beat_but_loss(trade: Dict, earnings_result: Dict) -> str:
    """
    Beat EPS but stock lost money. Try news research first; fall back to
    quantitative signals only when no relevant articles are found.
    """
    symbol        = trade["symbol"]
    earnings_date = trade.get("earnings_date") or trade.get("exit_date") or today_et().isoformat()

    news = fetch_post_earnings_news(symbol, earnings_date)
    if news:
        return (f"The model correctly identified the earnings beat but the stock sold off "
                f"post-announcement. {news}")

    # Quantitative fallback — only surfaces signals that are genuinely abnormal
    beat_pct = earnings_result.get("beat_pct") or 0
    momentum = trade.get("momentum_30d") or 0
    reasons  = []

    if momentum >= 15:
        reasons.append(
            f"a +{momentum:.0f}% pre-earnings run-up points to buy-the-rumour-sell-the-news"
        )
    if 0 < beat_pct < 3:
        reasons.append(
            f"the EPS beat was only +{beat_pct:.1f}%, likely below buy-side whisper numbers"
        )

    if reasons:
        return (f"The model correctly identified the earnings beat but the stock sold off — "
                f"quantitative signals suggest: {'; '.join(reasons)}.")

    return (f"The model correctly identified the earnings beat but the market reacted "
            f"negatively — possible guidance concerns or valuation resistance not captured "
            f"in the scoring model.")


def diagnose_miss_but_win(trade: Dict) -> str:
    """
    Missed EPS but stock went up. Try news research first; fall back to
    quantitative signals only when no relevant articles are found.
    """
    symbol        = trade["symbol"]
    earnings_date = trade.get("earnings_date") or trade.get("exit_date") or today_et().isoformat()

    news = fetch_post_earnings_news(symbol, earnings_date)
    if news:
        return (f"Despite the EPS miss the position was profitable. {news}")

    # Quantitative fallback
    momentum = trade.get("momentum_30d") or 0
    if momentum <= -10:
        return (f"Despite the EPS miss the position was profitable — the stock had already "
                f"sold off {abs(momentum):.0f}% pre-earnings, suggesting the miss was "
                f"largely priced in before the announcement.")

    return (f"Despite the EPS miss the position was profitable — possible forward guidance "
            f"upgrade or short squeeze not captured by the scoring model.")


def generate_analysis(trade: Dict, earnings_result: Dict) -> str:
    """Generate a 3-sentence analysis of the earnings call."""
    symbol   = trade["symbol"]
    pnl_pct  = trade.get("pnl_pct") or 0
    outcome  = "WIN" if (trade.get("pnl") or 0) > 0 else "LOSS"
    beat     = earnings_result.get("beat")
    eps_est  = earnings_result.get("eps_estimate")
    eps_rep  = earnings_result.get("eps_reported")
    beat_pct = earnings_result.get("beat_pct")
    trade_type = trade.get("trade_type", "call")

    # Sentence 1: earnings result
    if beat is True and eps_est is not None:
        s1 = f"{symbol} beat EPS estimates by {beat_pct:+.1f}% (reported ${eps_rep:.2f} vs ${eps_est:.2f} expected)."
    elif beat is False and eps_est is not None:
        s1 = f"{symbol} missed EPS estimates by {abs(beat_pct):.1f}% (reported ${eps_rep:.2f} vs ${eps_est:.2f} expected)."
    else:
        s1 = f"{symbol} earnings results not yet confirmed in data feed."

    # Sentence 2: trade outcome
    s2 = f"The Hitman v2 {trade_type} position returned {pnl_pct:+.1f}% ({outcome})."

    # Sentence 3: diagnosis
    if beat is None:
        s3 = "EPS result could not be verified — trade result logged for manual review."
    elif beat is True and outcome == "WIN":
        s3 = "Hitman v2 correctly identified the earnings beat and the market confirmed it."
    elif beat is False and outcome == "LOSS":
        s3 = "Both the earnings miss and the negative market reaction aligned with expectations."
    elif beat is True and outcome == "LOSS":
        s3 = diagnose_beat_but_loss(trade, earnings_result)
    elif beat is False and outcome == "WIN":
        s3 = diagnose_miss_but_win(trade)
    else:
        s3 = "Trade result flagged for manual review."

    return f"{s1} {s2} {s3}"


def update_earnings_log(today_closed: List[Dict]) -> None:
    """Append closed trades to earnings_calls_log.json with analysis."""
    if not today_closed:
        return
    log_entries = load_json("earnings_calls_log.json")
    if not isinstance(log_entries, list):
        log_entries = []
    existing_ids = {e.get("trade_id") for e in log_entries}

    accuracy = load_json("earnings_accuracy.json") or {
        "70-84%":  {"calls": 0, "correct": 0},
        "85-94%":  {"calls": 0, "correct": 0},
        "95-100%": {"calls": 0, "correct": 0},
    }

    # Re-check any existing entries where EPS data wasn't available yet
    for entry in log_entries:
        if entry.get("correct_call") is None and entry.get("beat") is None:
            earnings_result = fetch_earnings_result(entry["symbol"], entry.get("earnings_date", ""))
            beat = earnings_result.get("beat")
            if beat is not None:
                outcome = entry.get("outcome", "LOSS")
                correct_call = (beat is True and outcome == "WIN") or (beat is False and outcome == "LOSS")
                entry["beat"]         = beat
                entry["eps_estimate"] = earnings_result.get("eps_estimate")
                entry["eps_reported"] = earnings_result.get("eps_reported")
                entry["beat_pct"]     = earnings_result.get("beat_pct")
                entry["correct_call"] = correct_call
                entry["analysis"]     = generate_analysis(
                    {"symbol": entry["symbol"], "trade_type": entry.get("trade_type", "call"),
                     "pnl": entry.get("pnl", 0), "pnl_pct": entry.get("pnl_pct", 0),
                     "momentum_30d": entry.get("momentum_30d"), "earnings_date": entry.get("earnings_date"),
                     "exit_date": entry.get("exit_date")},
                    earnings_result,
                )
                log.info(f"{entry['symbol']}: EPS data now available — correct_call updated to {correct_call}")

    for trade in today_closed:
        trade_id = trade.get("id") or trade.get("trade_id")
        if trade_id in existing_ids:
            continue
        earnings_result = fetch_earnings_result(trade["symbol"], trade.get("earnings_date", ""))
        analysis = generate_analysis(trade, earnings_result)
        beat = earnings_result.get("beat")
        outcome = "WIN" if (trade.get("pnl") or 0) > 0 else "LOSS"
        correct_call = beat is not None and (
            (beat is True and outcome == "WIN") or (beat is False and outcome == "LOSS")
        )

        entry = {
            "trade_id":      trade_id,
            "symbol":        trade["symbol"],
            "company":       trade.get("company", trade["symbol"]),
            "entry_date":    trade.get("entry_date"),
            "earnings_date": trade.get("earnings_date"),
            "exit_date":     trade.get("exit_date"),
            "trade_type":    trade.get("trade_type", "call"),
            "beat_rate":     trade.get("beat_rate"),
            "pu5":           trade.get("pu5"),
            "momentum_30d":  trade.get("momentum_30d"),
            "iv_proxy":      trade.get("iv_proxy"),
            "long_contract": trade.get("long_contract"),
            "entry_price":   trade.get("entry_price"),
            "exit_price":    trade.get("exit_price"),
            "exit_value":    trade.get("exit_value"),
            "position_size": trade.get("position_size"),
            "pnl":           trade.get("pnl"),
            "pnl_pct":       trade.get("pnl_pct"),
            "outcome":       outcome,
            "eps_estimate":  earnings_result.get("eps_estimate"),
            "eps_reported":  earnings_result.get("eps_reported"),
            "beat_pct":      earnings_result.get("beat_pct"),
            "beat":          beat,
            "correct_call":  correct_call if beat is not None else None,
            "analysis":      analysis,
            "logged_at":     now_et().isoformat(),
        }
        log_entries.append(entry)

    # Recalculate accuracy from full log (idempotent — re-running never double-counts)
    accuracy = {"calls": 0, "correct": 0, "wins": 0}
    for e in log_entries:
        accuracy["calls"] += 1
        if e.get("correct_call"):
            accuracy["correct"] += 1
        if e.get("outcome") == "WIN":
            accuracy["wins"] += 1

    save_json("earnings_calls_log.json", log_entries)
    save_json("earnings_accuracy.json", accuracy)
    try:
        write_earnings_docx(log_entries, accuracy)
        log.info(f"Earnings Word doc updated: {EARNINGS_DOCX}")
    except Exception as e:
        log.warning(f"Could not write earnings Word doc: {e}")
    log.info(f"Earnings log updated: {len(today_closed)} new entries")


def write_earnings_docx(log_entries: List[Dict], accuracy: Dict) -> None:
    """Regenerate the full earnings_calls_log.docx from the log entries."""
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        log.warning("python-docx not installed — skipping .docx generation")
        return

    doc = Document()

    # Title
    title = doc.add_heading("Earnings Calls Log", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    srun = sub.add_run(f"Generated {now_et().strftime('%Y-%m-%d %H:%M ET')}  |  {len(log_entries)} calls logged")
    srun.italic = True
    srun.font.size = Pt(10)

    # Accuracy summary
    doc.add_heading("Hitman v2 Accuracy Summary", level=1)
    calls   = accuracy.get("calls", 0)
    correct = accuracy.get("correct", 0)
    wins    = accuracy.get("wins", 0)
    acc_pct = (correct / calls * 100) if calls else 0
    win_pct = (wins / calls * 100) if calls else 0
    table = doc.add_table(rows=1, cols=3)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    hdr[0].text = "Total Calls"
    hdr[1].text = "Direction Accuracy"
    hdr[2].text = "Win Rate (P&L)"
    row = table.add_row().cells
    row[0].text = str(calls)
    row[1].text = f"{acc_pct:.0f}%" if calls else "—"
    row[2].text = f"{win_pct:.0f}%" if calls else "—"

    # Individual calls, newest first
    doc.add_heading("Call Details", level=1)
    sorted_entries = sorted(
        log_entries,
        key=lambda e: e.get("exit_date") or e.get("logged_at") or "",
        reverse=True,
    )
    for entry in sorted_entries:
        symbol = entry.get("symbol", "?")
        company = entry.get("company", symbol)
        outcome = entry.get("outcome", "?")
        pnl = entry.get("pnl") or 0
        pnl_pct = entry.get("pnl_pct") or 0

        h = doc.add_heading(level=2)
        hrun = h.add_run(f"{symbol} — {company}  ({outcome})")
        hrun.font.color.rgb = RGBColor(0x2E, 0x7D, 0x32) if outcome == "WIN" else RGBColor(0xC6, 0x28, 0x28)

        meta = doc.add_paragraph()
        meta.add_run("Earnings date: ").bold = True
        meta.add_run(f"{entry.get('earnings_date', '—')}   ")
        meta.add_run("Exit: ").bold = True
        meta.add_run(f"{entry.get('exit_date', '—')}   ")
        meta.add_run("Trade type: ").bold = True
        meta.add_run(f"{entry.get('trade_type', '?')}   ")
        meta.add_run("IV proxy: ").bold = True
        meta.add_run(f"{entry.get('iv_proxy', 0):.0%}" if entry.get('iv_proxy') else "—")

        prices = doc.add_paragraph()
        prices.add_run("Entry cost: ").bold = True
        prices.add_run(f"${entry.get('position_size', 0):,.0f}   ")
        prices.add_run("Exit value: ").bold = True
        prices.add_run(f"${entry.get('exit_value', entry.get('exit_price', 0)):,.2f}   ")
        prices.add_run("P&L: ").bold = True
        prices.add_run(f"${pnl:+,.2f} ({pnl_pct:+.1f}%)")

        eps_est = entry.get("eps_estimate")
        eps_rep = entry.get("eps_reported")
        beat_pct = entry.get("beat_pct")
        beat = entry.get("beat")
        eps_p = doc.add_paragraph()
        eps_p.add_run("EPS: ").bold = True
        if eps_est is not None and eps_rep is not None:
            label = "BEAT" if beat else ("MISS" if beat is False else "—")
            eps_p.add_run(f"${eps_rep:.2f} reported vs ${eps_est:.2f} est ({beat_pct:+.1f}%) — {label}")
        else:
            eps_p.add_run("not available")

        correct_call = entry.get("correct_call")
        call_p = doc.add_paragraph()
        call_p.add_run("Call quality: ").bold = True
        if correct_call is True:
            call_p.add_run("Algorithm aced the call ✓")
        elif correct_call is False:
            call_p.add_run("Algorithm missed the call ✗")
        else:
            call_p.add_run("Unverified")

        ana = doc.add_paragraph()
        ana.add_run("Analysis: ").bold = True
        ana.add_run(entry.get("analysis", ""))

        doc.add_paragraph("─" * 60).alignment = WD_ALIGN_PARAGRAPH.CENTER

    DATA_DIR.mkdir(exist_ok=True)
    doc.save(str(EARNINGS_DOCX))


def run_reporter(report_date: Optional[str] = None):
    """
    Generate the daily performance report.
    Pass report_date (YYYY-MM-DD) to generate a report for a past day.
    Defaults to today (ET).
    """
    log.info("=" * 60)
    log.info("REPORTER: Generating daily performance report")
    log.info("=" * 60)

    if report_date:
        target_date = report_date
        log.info(f"Running for historical date: {target_date}")
    else:
        if not is_weekday():
            log.info("Not a weekday — skipping report")
            return
        target_date = today_et().isoformat()

    trades_history = load_json("trades_history.json") or []
    positions = load_json("positions.json") or {}
    budget = load_json("budget_tracker.json") or {}

    today_trades = [t for t in trades_history if t.get("entry_date") == target_date]
    today_closed = [
        t for t in trades_history
        if t.get("exit_date") == target_date and t.get("status") == "closed"
    ]
    all_closed = get_all_closed(trades_history)

    update_earnings_log(today_closed)
    text = build_text_report(today_trades, today_closed, all_closed, positions, budget)
    # Patch the date header in the text to show the correct report date
    text = text.replace(today_et().isoformat(), target_date, 1)
    html = build_html_report(text, today_closed)
    log.info("\n" + text)
    daily_dir = REPORTS_DIR / "Daily Model Movement"
    daily_dir.mkdir(parents=True, exist_ok=True)
    report_file = daily_dir / f"report_{target_date}.html"

    report_file.write_text(html, encoding="utf-8")
    log.info(f"Report saved: {report_file}")
    daily_pnl = positions.get("daily_pnl", 0)
    win_rate = calc_win_rate(all_closed)
    send_email(
        f"Daily Report ({target_date}): P&L ${daily_pnl:+.2f} | Win Rate {win_rate:.0f}%",
        text,
        html,
    )


if __name__ == "__main__":
    import sys
    date_arg = sys.argv[1] if len(sys.argv) > 1 else None
    run_reporter(report_date=date_arg)
